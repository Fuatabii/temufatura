import streamlit as st
import pandas as pd
import fitz
import re
from io import BytesIO
from zipfile import ZipFile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import unidecode
import msoffcrypto
import xlsxwriter
from docx.enum.section import WD_ORIENT


st.set_page_config(page_title="📦 Proforma Invoice Oluşturucu", layout="wide")
st.title("📄 AWB + Şifreli/Normal Excel'den Proforma Invoice")

EXCEL_PASSWORD = "PGgf!7p*Bm"

def read_excel_with_optional_password(uploaded_file):
    try:
        return pd.read_excel(uploaded_file, engine="openpyxl", dtype=str)
    except Exception:
        try:
            uploaded_file.seek(0)
            office_file = msoffcrypto.OfficeFile(uploaded_file)
            office_file.load_key(password=EXCEL_PASSWORD)
            decrypted = BytesIO()
            office_file.decrypt(decrypted)
            return pd.read_excel(decrypted, engine="openpyxl", dtype=str)
        except Exception as e:
            st.warning(f"{uploaded_file.name} şifreli ama açılamadı: {e}")
            return None

def extract_awb_data_multi(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    results = []

    for page in doc:
        text = page.get_text()
        awb_matches = re.findall(r"716-\d{8}", text)
        for awb_number in set(awb_matches):
            packages = re.search(r"TOTAL\s*:\s*\((\d+)\)\s*PACKAGES", text)
            gross = re.search(r"(\d+)\s+K\s+Q\s+\d+\s+", text)
            volume = re.search(r"VOL\s*[:：]?\s*([\d.]+)CBM", text, re.IGNORECASE)
            dims = re.search(r"DIM[:：]?\s*([0-9* /]+)", text)

            sender_name_match = re.search(r"(SMART-TRANS LOGISTIC CHENGDU LTD.*?)\s{2,}", text)
            sender_address_match = re.search(
                r"SMART-TRANS LOGISTIC CHENGDU LTD\s+(NO\..+?CHINA[\uff08(]?\d{6}[\uff09)]?)",
                text,
                re.DOTALL
            )

            sender_name = sender_name_match.group(1).strip() if sender_name_match else ""
            sender_address = sender_address_match.group(1).replace("\n", " ").strip() if sender_address_match else ""

            results.append({
                "AWB Number": awb_number,
                "Packages": int(packages.group(1)) if packages else None,
                "Gross Weight": int(gross.group(1)) if gross else None,
                "Volume": float(volume.group(1)) if volume else None,
                "Dimensions": dims.group(1).strip() if dims else None,
                "SenderName": sender_name,
                "SenderAddress": sender_address,
                "Text": text
            })
    return results

def create_proforma(df, awb_info):
    doc = Document()

    # LANDSCAPE ayarı ve marginler
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.1)
    section.right_margin = Inches(0.1)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.1)

    # Başlık ve taraf bilgileri
    doc.add_paragraph("PROFORMA INVOICE", style='Heading 1')

    doc.add_paragraph(
        f"Sender\n{awb_info['SenderName']}\n{awb_info['SenderAddress']}"
    )

    doc.add_paragraph(
        "Receiver\nBOX NOW SA - Tatoiou 96, Acharne 13672, Athens Greece.\n"
        "Local Warehouse: SkyServe (Supervision Warehouse), Spata, 19004\n"
        "Customs code GR000304"
    )

    doc.add_paragraph(f"Tarih / Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")

    # Tablo oluştur
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False

    col_widths = [Inches(0.9), Inches(4.1), Inches(4.6), Inches(0.7), Inches(0.5)]
    headers = ['TrackingNumber', 'GoodsDescription', 'HS Code', 'Quantity', 'Total Value (EUR)']

    # Başlık satırı
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(header)
        run.font.size = Pt(7)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Sayısal dönüşümler
    df['InvoiceAmountTotal'] = pd.to_numeric(df['InvoiceAmountTotal'], errors='coerce').fillna(0)
    df['ItemPackageQuantity'] = pd.to_numeric(df['ItemPackageQuantity'], errors='coerce').fillna(0)

    # Gruplama
    df_grouped = df.groupby('TrackingNumber').agg({
        'GoodsDescription': lambda x: ', '.join(sorted(set(x))),
        'CommodityTaric': lambda x: ', '.join(sorted(set(x))),
        'ItemPackageQuantity': 'sum',
        'InvoiceAmountTotal': 'first'
    }).reset_index()

    total_value = 0

    for _, row in df_grouped.iterrows():
        row_cells = table.add_row().cells
        values = [
            str(row['TrackingNumber']),
            str(row['GoodsDescription']),
            str(row['CommodityTaric']),
            str(int(row['ItemPackageQuantity'])),
            f"{float(row['InvoiceAmountTotal']):.2f}"
        ]
        total_value += float(row['InvoiceAmountTotal'])

        for i, val in enumerate(values):
            cell = row_cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if i == 1 else WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(val)
            run.font.size = Pt(6.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Alt bilgi bloğu
    doc.add_paragraph("")

    summary = doc.add_paragraph()
    summary.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    total_kap_text = f"TOTAL KAP: {awb_info.get('Packages', 0)}      {awb_info.get('Gross Weight', 0)} kg"
    mawb_text = f"MAWB NO : {awb_info.get('AWB Number', '')}"
    total_pcs_text = f"TOTAL PCS: {int(df['ItemPackageQuantity'].sum())}"
    hawb_text = f"HAWB: {df['TrackingNumber'].nunique()}"

    for line in [total_kap_text, mawb_text, total_pcs_text, hawb_text]:
        run = summary.add_run(line + "\n")
        run.bold = True
        run.font.size = Pt(6.5)

    total_line = doc.add_paragraph()
    total_line.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = total_line.add_run(f"TOPLAM     {total_value:,.2f} €")
    run.bold = True
    run.font.size = Pt(6.5)

    return doc

def create_summary_excel(summary_list):
    output = BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet("Özet Rapor")

    headers = ["MAWB No", "Kap Adedi", "KG", "HAWB", "FATURA TOPLAM"]
    header_format = workbook.add_format({'bold': True, 'bg_color': '#D9D9D9', 'border': 1})
    number_format = workbook.add_format({'num_format': '#,##0.00', 'border': 1})
    int_format = workbook.add_format({'num_format': '0', 'border': 1})

    for col, header in enumerate(headers):
        worksheet.write(0, col, header, header_format)

    for row_idx, row in enumerate(summary_list, start=1):
        worksheet.write(row_idx, 0, row['MAWB No'], int_format)
        worksheet.write(row_idx, 1, row['Kap Adedi'], int_format)
        worksheet.write(row_idx, 2, row['KG'], int_format)
        worksheet.write(row_idx, 3, row['HAWB'], int_format)
        worksheet.write(row_idx, 4, row['FATURA TOPLAM'], number_format)

    worksheet.set_column("A:A", 18)
    worksheet.set_column("B:E", 15)
    workbook.close()
    output.seek(0)
    return output

uploaded_pdfs = st.file_uploader("📄 AWB PDF Dosyaları", type=["pdf"], accept_multiple_files=True)
uploaded_excels = st.file_uploader("📊 Excel Manifestoları (.xlsx)", type=["xlsx"], accept_multiple_files=True)

if uploaded_pdfs and uploaded_excels and st.button("🔄 Proformaları Oluştur ve ZIP Olarak İndir"):
    with st.spinner("Faturalar hazırlanıyor..."):

        progress_bar = st.progress(0)
        status_text = st.empty()
        output_files = []
        summary_list = []

        all_awbs = []
        awb_mapping = {}

        for pdf in uploaded_pdfs:
            awbs = extract_awb_data_multi(pdf)
            all_awbs.extend(awbs)
            awb_mapping[pdf.name] = awbs

        total_awbs = len(all_awbs)
        current_awb = 0

        for pdf in uploaded_pdfs:
            awb_list = awb_mapping[pdf.name]

            for awb_info in awb_list:
                current_awb += 1
                status_text.text(f"📦 İşleniyor: {awb_info['AWB Number']}  ({current_awb}/{total_awbs})")
                progress_bar.progress(current_awb / total_awbs)

                matched_df = None
                for excel in uploaded_excels:
                    df = read_excel_with_optional_password(excel)
                    if df is None:
                        continue
                    if 'BX-M-N741' in df.columns and awb_info['AWB Number'] in df['BX-M-N741'].astype(str).values:
                        matched_df = df.copy()
                        break

                if matched_df is not None:
                    for col in ['TrackingNumber', 'GoodsDescription', 'CommodityTaric', 'ItemPackageQuantity', 'InvoiceAmountTotal']:
                        if col in matched_df.columns:
                            matched_df[col] = matched_df[col].astype(str).apply(lambda x: unidecode.unidecode(x))

                    doc = create_proforma(matched_df, awb_info)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    filename = f"{awb_info['AWB Number'].replace('-', '')}.docx"
                    output_files.append((filename, buffer.read()))

                    merged = matched_df.groupby("TrackingNumber").agg({"InvoiceAmountTotal": "first"}).reset_index()
                    merged["InvoiceAmountTotal"] = pd.to_numeric(merged["InvoiceAmountTotal"], errors="coerce").fillna(0)

                    summary_list.append({
                        "MAWB No": awb_info["AWB Number"],
                        "Kap Adedi": awb_info["Packages"] or 0,
                        "KG": awb_info["Gross Weight"] or 0,
                        "HAWB": matched_df["TrackingNumber"].nunique(),
                        "FATURA TOPLAM": merged["InvoiceAmountTotal"].sum()
                    })
                else:
                    st.warning(f"❌ Eşleşen Excel bulunamadı: {awb_info['AWB Number']}")

        if output_files:
            summary_excel = create_summary_excel(summary_list)

            zip_buffer = BytesIO()
            with ZipFile(zip_buffer, "w") as zipf:
                for name, data in output_files:
                    zipf.writestr(name, data)
                zipf.writestr("ozet_rapor.xlsx", summary_excel.read())
            zip_buffer.seek(0)

            status_text.text("✅ Tüm işlemler tamamlandı.")
            progress_bar.progress(1.0)
            st.success(f"{len(output_files)} adet Proforma Invoice ve Özet Rapor ZIP içinde hazır!")

            st.markdown("### 📦 İndirme")
            st.download_button(
                label="📥 ZIP (Faturalar + Özet Rapor)",
                data=zip_buffer,
                file_name="proforma_invoices.zip",
                mime="application/zip"
            )
        else:
            status_text.text("⚠️ Hiçbir dosya işlenemedi.")
            progress_bar.empty()
